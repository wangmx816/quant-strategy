#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成 TASK3 Word 报告：双均线策略理论与回测分析。"""

from __future__ import annotations

import os
import sys
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(ROOT))

from src.backtest import BacktestConfig, run_backtest, run_multi_experiments
from src.data_fetch import DATA_DIR, STOCKS, fetch_all
from src.diagnostics import diagnose, format_diagnostic_text

CHART_DIR = ROOT / "output" / "charts"
DOC_PATH = ROOT / f"{os.environ.get('STUDENT_NAME', '姓名')}TASK3.docx"
REPO_URL = "https://github.com/wangmx816/quant-strategy"
PAGES_URL = "https://wangmx816.github.io/quant-strategy/"


def set_run_font(run, size_pt: float = 10.5, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def add_paragraph(doc: Document, text: str, first_line_indent: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, bold=True)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run)


def load_data() -> dict[str, pd.DataFrame]:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    if not (DATA_DIR / "002202_daily.csv").exists():
        return fetch_all("qfq")
    data = {}
    for symbol in STOCKS:
        df = pd.read_csv(DATA_DIR / f"{symbol}_daily.csv", parse_dates=["trade_date"])
        data[symbol] = df.sort_values("trade_date").reset_index(drop=True)
    return data


def plot_goldwind_strategy(df: pd.DataFrame, out: Path) -> dict:
    cfg = BacktestConfig(short_window=5, long_window=15)
    res = run_backtest(df, cfg)
    d, eq = res["data"], res["equity"]
    plt.rcParams["font.sans-serif"] = ["SimSun", "Microsoft YaHei"]
    plt.rcParams["axes.unicode_minus"] = False

    fig, axes = plt.subplots(2, 1, figsize=(10, 7), gridspec_kw={"height_ratios": [2, 1]})
    ax1, ax2 = axes
    ax1.plot(d["trade_date"], d["close"], color="#f97316", linewidth=1.2, label="收盘价")
    ax1.plot(d["trade_date"], d["sma_short"], color="#38bdf8", linewidth=1, label="SMA(5)")
    ax1.plot(d["trade_date"], d["sma_long"], color="#1d4ed8", linewidth=1, label="SMA(15)")
    buys = d[d["buy_signal"] == 1]
    sells = d[d["sell_signal"] == 1]
    ax1.scatter(buys["trade_date"], buys["close"], marker="^", color="#dc2626", s=60, label="买入", zorder=5)
    ax1.scatter(sells["trade_date"], sells["close"], marker="v", color="#16a34a", s=60, label="卖出", zorder=5)
    ax1.set_title("图1  金风科技双均线策略：价格、均线与交易信号", fontsize=13, pad=10)
    ax1.legend(loc="upper left", fontsize=9)
    ax1.grid(True, linestyle="--", alpha=0.3)

    ax2.plot(eq["trade_date"], eq["net_value"], color="#dc2626", label="策略净值")
    ax2.plot(d["trade_date"], d["benchmark_nv"], color="#94a3b8", label="买入持有")
    ax2.set_title("策略净值 vs 买入持有基准", fontsize=11)
    ax2.legend(loc="upper left", fontsize=9)
    ax2.grid(True, linestyle="--", alpha=0.3)
    fig.autofmt_xdate()
    fig.tight_layout()
    fig.savefig(out, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return res["metrics"]


def plot_multi_stock_compare(exp: pd.DataFrame, out: Path) -> None:
    base = exp[(exp["short_ma"] == 5) & (exp["long_ma"] == 15)].copy()
    base["name"] = base["symbol"].map(lambda s: STOCKS[s]["name"])
    plt.rcParams["font.sans-serif"] = ["SimSun", "Microsoft YaHei"]
    plt.rcParams["axes.unicode_minus"] = False

    fig, axes = plt.subplots(2, 2, figsize=(11, 8))
    names = base["name"].tolist()
    x = np.arange(len(names))

    for ax, col, title in zip(
        axes.flat,
        ["annualized_return", "sharpe_ratio", "max_drawdown", "excess_return"],
        ["年化收益率 (%)", "夏普比率", "最大回撤 (%)", "超额收益 (%)"],
    ):
        vals = base[col].values
        colors = ["#dc2626" if v >= 0 else "#16a34a" for v in vals] if col != "max_drawdown" else ["#16a34a"] * len(vals)
        if col == "max_drawdown":
            vals = np.abs(vals)
        bars = ax.bar(x, vals, color=colors, width=0.55)
        ax.set_title(title, fontsize=11)
        ax.set_xticks(x)
        ax.set_xticklabels(names, fontsize=9)
        ax.grid(axis="y", linestyle="--", alpha=0.3)
        for b, v in zip(bars, vals):
            ax.text(b.get_x() + b.get_width() / 2, b.get_height(), f"{v:.1f}", ha="center", va="bottom", fontsize=8)

    fig.suptitle("图2  五只股票双均线 SMA(5/15) 回测指标对比", fontsize=13, y=1.02)
    fig.tight_layout()
    fig.savefig(out, dpi=180, bbox_inches="tight")
    plt.close(fig)


def plot_ma_param_sensitivity(exp: pd.DataFrame, symbol: str, out: Path) -> None:
    sub = exp[exp["symbol"] == symbol].copy()
    sub["label"] = sub.apply(lambda r: f"SMA({int(r['short_ma'])}/{int(r['long_ma'])})", axis=1)
    plt.rcParams["font.sans-serif"] = ["SimSun", "Microsoft YaHei"]
    plt.rcParams["axes.unicode_minus"] = False
    fig, ax = plt.subplots(figsize=(9, 4.5))
    x = np.arange(len(sub))
    colors = ["#dc2626" if v >= 0 else "#16a34a" for v in sub["annualized_return"]]
    ax.bar(x, sub["annualized_return"], color=colors, width=0.5)
    ax.set_xticks(x)
    ax.set_xticklabels(sub["label"])
    ax.set_ylabel("年化收益率 (%)")
    ax.set_title(f"图3  {STOCKS[symbol]['name']}不同均线参数下的年化收益", fontsize=12)
    ax.grid(axis="y", linestyle="--", alpha=0.3)
    fig.tight_layout()
    fig.savefig(out, dpi=180, bbox_inches="tight")
    plt.close(fig)


def build_report(stock_data: dict[str, pd.DataFrame]) -> None:
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    gw_metrics = plot_goldwind_strategy(stock_data["002202"], CHART_DIR / "fig1_goldwind_ma.png")
    exp = run_multi_experiments(stock_data)
    plot_multi_stock_compare(exp, CHART_DIR / "fig2_multi_stock.png")
    plot_ma_param_sensitivity(exp, "002202", CHART_DIR / "fig3_ma_sensitivity.png")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    add_heading(doc, "TASK3  双均线策略回测与量化指标分析")
    add_paragraph(doc, "课程：量化交易基础 · 作业三：策略回测与看板实现")

    add_heading(doc, "一、双均线策略基本概念")
    add_paragraph(
        doc,
        "双均线策略（Dual Moving Average Crossover）是最经典的趋势跟踪策略之一。"
        "其核心思想是用两条不同周期的简单移动平均线（SMA）刻画短中期价格趋势："
        "短期均线对价格变化更敏感，长期均线更平滑。当短期均线上穿长期均线时，"
        "称为金叉（Golden Cross），通常被视为买入信号，意味着短期动量转强；"
        "当短期均线下穿长期均线时，称为死叉（Death Cross），视为卖出信号，"
        "意味着短期动量转弱。策略假设价格趋势具有持续性，通过跟随趋势获取收益。"
        "但该策略在震荡市中容易产生频繁交叉，导致“来回打脸”和交易成本侵蚀。",
    )
    add_paragraph(
        doc,
        "本任务默认采用 SMA(5) 与 SMA(15) 组合：5 日反映近一周价格均值，"
        "15 日反映近三周趋势。信号规则为：SMA(5) > SMA(15) 时持有多头，"
        "否则空仓。实际回测中还需考虑手续费（本任务按万三 0.03%）、"
        "滑点（万一 0.01%）及全仓/固定比例仓位管理，否则回测结果会过于乐观。",
    )

    add_heading(doc, "二、量化策略效果基础指标")
    add_paragraph(
        doc,
        "累计回报（Cumulative Return）衡量策略在整个回测区间内总体的盈亏幅度，"
        "计算公式为（期末净值 − 1）× 100%。它直观反映“一共赚了多少”，"
        "但未考虑时间长度，不便跨区间比较。",
    )
    add_paragraph(
        doc,
        "年化收益率（Annualized Return）将累计收益换算为按年计的标准化收益，"
        "便于不同回测区间之间的比较。本报告采用复利年化公式，"
        "基于日度净值序列估算。",
    )
    add_paragraph(
        doc,
        "最大回撤（Maximum Drawdown, MDD）是净值从历史最高点回落至最低点的最大幅度，"
        "衡量策略可能面临的最极端账面损失。MDD 越小，说明策略风险控制越好；"
        "投资者需关注 MDD 是否在可承受范围内。",
    )
    add_paragraph(
        doc,
        "夏普比率（Sharpe Ratio）衡量单位风险所获得的超额收益，"
        "计算公式为（策略平均超额日收益 / 日收益标准差）× √252。"
        "夏普比率越高，说明风险调整后收益越好；通常大于 1 视为较好，"
        "0.5~1 为中等，低于 0.5 则风险收益比偏弱。",
    )
    add_paragraph(
        doc,
        "此外，本报告还统计胜率（盈利交易占比）、盈亏比（平均盈利/平均亏损）"
        "及相对买入持有基准的超额收益，用于综合评估策略是否真正创造价值。",
    )

    add_heading(doc, "三、数据加载与基础诊断")
    add_paragraph(
        doc,
        "本任务选取金风科技（002202.SZ）、三一重工（600031.SH）、徐工机械（000425.SZ）、"
        "安彩高科（600207.SH）、智慧农业（000816.SZ）五只 A 股，"
        "通过东方财富 API 获取近一年前复权日线 OHLCV 数据。"
        "说明：题目中“安彩高科666207”应为“600207”之误，本报告按 600207.SH 处理。",
    )
    for symbol, df in stock_data.items():
        diag = diagnose(df, symbol)
        add_paragraph(doc, format_diagnostic_text(diag))

    add_heading(doc, "四、Python 编程实现")
    add_paragraph(
        doc,
        "项目结构：quant-strategy/ 包含 src/data_fetch.py（数据获取）、"
        "src/diagnostics.py（数据诊断）、src/backtest.py（双均线信号与回测）、"
        "build_dashboard.py（HTML 看板生成）及 generate_task3_report.py（Word 报告）。"
        "核心流程为：① 加载 CSV 日线数据；② 计算 SMA 短/长均线；"
        "③ 识别金叉/死叉生成买卖信号；④ 模拟交易并扣除手续费与滑点；"
        "⑤ 计算累计回报、年化收益、MDD、夏普等指标；⑥ 输出 Chart.js 交互看板。",
    )
    add_paragraph(
        doc,
        f"GitHub 仓库：{REPO_URL}；在线看板：{PAGES_URL}。"
        "看板包含侧边栏参数区、KPI 卡片、策略净值图、回撤图、价格均线买卖点图"
        "及多股票多参数对比图，与课程示例界面要素一致。",
    )

    add_heading(doc, "五、金风科技回测结果与图表解读")
    m = gw_metrics
    add_paragraph(
        doc,
        f"以金风科技 SMA(5/15) 为例，回测区间年化收益率 {m['annualized_return']:.2f}%，"
        f"累计回报 {m['cumulative_return']:.2f}%，最大回撤 {m['max_drawdown']:.2f}%"
        f"（{m['max_dd_start']} 至 {m['max_dd_end']}），夏普比率 {m['sharpe_ratio']:.3f}，"
        f"共 {m['trade_count']} 笔交易，胜率 {m['win_rate']:.1f}%，"
        f"相对买入持有超额收益 {m['excess_return']:.2f}%。",
    )
    doc.add_picture(str(CHART_DIR / "fig1_goldwind_ma.png"), width=Cm(15))
    add_caption(doc, "图1  金风科技双均线策略：价格、均线与交易信号")
    add_paragraph(
        doc,
        "图1上半部分展示收盘价与 SMA(5)、SMA(15) 的相对位置，"
        "红色上三角标记金叉买入点，绿色下三角标记死叉卖出点。"
        "可见 2025 年下半年至 2026 年初金风科技出现明显趋势上涨，"
        "双均线策略在趋势段能较好跟随；但在震荡回落阶段，"
        "均线反复交叉可能导致滞后止损。下半部分对比策略净值与买入持有："
        "若策略曲线低于基准，说明简单双均线未能跑赢被动持有。",
    )

    add_heading(doc, "六、多股票与多参数对比实验")
    doc.add_picture(str(CHART_DIR / "fig2_multi_stock.png"), width=Cm(15))
    add_caption(doc, "图2  五只股票双均线 SMA(5/15) 回测指标对比")
    base = exp[(exp["short_ma"] == 5) & (exp["long_ma"] == 15)]
    summary_lines = []
    for _, r in base.iterrows():
        summary_lines.append(
            f"{STOCKS[r['symbol']]['name']}：年化 {r['annualized_return']:.1f}%，"
            f"夏普 {r['sharpe_ratio']:.2f}，MDD {r['max_drawdown']:.1f}%，"
            f"超额 {r['excess_return']:.1f}%"
        )
    add_paragraph(doc, "SMA(5/15) 下五只股票表现：" + "；".join(summary_lines) + "。")
    add_paragraph(
        doc,
        "图2显示，不同标的在同一策略下的表现差异显著，说明双均线策略具有明显的"
        "“品种选择效应”：趋势性强、波动适中的股票更适合；"
        "持续单边下跌或剧烈震荡的标的，策略往往跑输买入持有。"
        "工程机械板块（三一、徐工）与农业股（智慧农业）在本区间表现需结合行业周期解读。",
    )

    doc.add_picture(str(CHART_DIR / "fig3_ma_sensitivity.png"), width=Cm(14))
    add_caption(doc, "图3  金风科技不同均线参数下的年化收益")
    add_paragraph(
        doc,
        "图3对比 SMA(5/15)、(5/20)、(5/35)、(10/30) 四组参数。"
        "长期均线周期越大，信号越滞后但越平滑，交易次数减少、手续费影响降低；"
        "短期均线越小，反应越快但假信号增多。实践中需在灵敏度与稳定性之间权衡。",
    )

    add_heading(doc, "七、双均线策略适用场景与应用心得")
    add_paragraph(
        doc,
        "适用场景：（1）趋势明确的中长期行情，如板块景气上行或个股基本面改善驱动的上涨；"
        "（2）流动性好、交易成本占比较低的标的；（3）作为更大策略体系的组成部分，"
        "与止损、仓位管理、基本面过滤结合使用。",
    )
    add_paragraph(
        doc,
        "不适用场景：（1）横盘震荡市，均线频繁交叉导致连续小额亏损；"
        "（2）跳空缺口多、政策扰动大的小盘股；（3）忽视交易成本的高频化改造——"
        "如截图所示，不考虑成本的“漂亮回测”与扣除成本后的实盘可能相差悬殊。",
    )
    add_paragraph(
        doc,
        "应用心得：① 务必使用前复权数据，避免除权造成虚假信号；"
        "② 必须扣除手续费与滑点，万三+万一已属 A 股较低水平；"
        "③ 不要只看累计收益，需同时关注 MDD 与夏普；"
        "④ 与买入持有基准对比，策略若不能产生稳定超额收益则意义有限；"
        "⑤ 参数不宜过度优化，样本外验证才能检验策略鲁棒性。",
    )

    add_heading(doc, "八、项目链接")
    add_paragraph(
        doc,
        f"完整代码与交互看板已推送至 GitHub：{REPO_URL}。"
        f"在线访问地址：{PAGES_URL}。"
        "本地运行：pip install -r requirements.txt，"
        "python build_dashboard.py 生成看板，python generate_task3_report.py 生成报告。",
    )

    doc.save(DOC_PATH)
    print(f"doc={DOC_PATH}")


def main() -> None:
    stock_data = load_data()
    build_report(stock_data)


if __name__ == "__main__":
    main()
